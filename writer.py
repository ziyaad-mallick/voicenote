"""
Writes formatted notes to disk as Markdown (with Obsidian-compatible frontmatter)
and optionally as .docx.
"""

from __future__ import annotations
import re
from datetime import datetime
from pathlib import Path


def slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "-", text)
    return text[:60]


def write_markdown(note: dict, notes_dir: Path) -> Path:
    category = note.get("category", "Personal")
    date_str = datetime.now().strftime("%Y-%m-%d")
    time_str = datetime.now().strftime("%H:%M")
    title_slug = slugify(note.get("title", "untitled"))

    out_dir = notes_dir / category
    out_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{date_str}_{title_slug}.md"
    path = out_dir / filename

    tags_yaml = "\n".join(f"  - {t}" for t in note.get("tags", []))
    reminders_yaml = ""
    for r in note.get("reminders", []):
        reminders_yaml += f"\n  - text: \"{r['text']}\"\n    datetime: \"{r['datetime']}\""

    frontmatter = f"""\
---
title: "{note['title']}"
date: {date_str} {time_str}
category: {category}
tags:
{tags_yaml if tags_yaml else "  []"}
summary: "{note.get('summary', '').replace('"', "'")}"
reminders:{reminders_yaml if reminders_yaml else " []"}
---

"""
    content = frontmatter + f"# {note['title']}\n\n"
    content += f"> {note.get('summary', '')}\n\n"
    content += note.get("body", "")

    path.write_text(content, encoding="utf-8")
    return path


def write_docx(note: dict, notes_dir: Path) -> Path:
    from docx import Document
    from docx.shared import Pt
    import re as _re

    category = note.get("category", "Personal")
    date_str = datetime.now().strftime("%Y-%m-%d")
    title_slug = slugify(note.get("title", "untitled"))

    out_dir = notes_dir / category
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{date_str}_{title_slug}.docx"

    doc = Document()
    doc.add_heading(note["title"], 0)

    p = doc.add_paragraph()
    p.add_run(f"Date: {date_str}  |  Category: {category}").bold = True
    doc.add_paragraph(note.get("summary", ""))
    doc.add_heading("Notes", level=1)

    # Render markdown-ish body into docx (basic: strip ### and ** **)
    body = note.get("body", "")
    for line in body.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            clean = _re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            doc.add_paragraph(clean)

    if note.get("reminders"):
        doc.add_heading("Reminders", level=1)
        for r in note["reminders"]:
            doc.add_paragraph(f"{r['text']} — {r['datetime']}", style="List Bullet")

    doc.save(str(path))
    return path
