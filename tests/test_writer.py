"""writer.py: pure formatting/filesystem logic — no mocking needed.

`slugify` and the Markdown frontmatter assembly are pure string logic.
`write_markdown` / `write_docx` touch the filesystem, but only through
`tmp_path`, and `datetime.now()` is frozen so filenames are deterministic.
"""
from datetime import datetime as _real_datetime

import pytest

import writer


class _FrozenDateTime(_real_datetime):
    @classmethod
    def now(cls, tz=None):
        return cls(2026, 1, 15, 9, 30)


@pytest.fixture
def frozen_time(monkeypatch):
    monkeypatch.setattr(writer, "datetime", _FrozenDateTime)


def _note(**overrides):
    note = {
        "title": "Buy groceries",
        "category": "Personal",
        "summary": "Need milk and eggs.",
        "body": "Buy milk and eggs from the store.",
        "tags": ["groceries", "errand"],
        "reminders": [],
    }
    note.update(overrides)
    return note


# -- slugify ------------------------------------------------------------

def test_slugify_lowercases_and_hyphenates_spaces():
    assert writer.slugify("Buy Groceries Today") == "buy-groceries-today"


def test_slugify_strips_punctuation():
    assert writer.slugify("Call mum!! (urgent)") == "call-mum-urgent"


def test_slugify_collapses_underscores_and_repeated_hyphens():
    assert writer.slugify("a__b---c") == "a-b-c"


def test_slugify_truncates_to_sixty_chars():
    assert len(writer.slugify("x " * 100)) == 60


def test_slugify_of_empty_string_is_empty():
    assert writer.slugify("") == ""


# -- write_markdown -------------------------------------------------------

def test_write_markdown_creates_a_category_subdirectory(tmp_path, frozen_time):
    path = writer.write_markdown(_note(category="Ideas"), tmp_path)
    assert path.parent == tmp_path / "Ideas"
    assert path.exists()


def test_write_markdown_filename_uses_frozen_date_and_slug(tmp_path, frozen_time):
    path = writer.write_markdown(_note(), tmp_path)
    assert path.name == "2026-01-15_buy-groceries.md"


def test_write_markdown_body_contains_title_summary_and_body(tmp_path, frozen_time):
    path = writer.write_markdown(_note(), tmp_path)
    content = path.read_text(encoding="utf-8")
    assert "# Buy groceries" in content
    assert "> Need milk and eggs." in content
    assert "Buy milk and eggs from the store." in content


def test_write_markdown_frontmatter_lists_each_tag(tmp_path, frozen_time):
    path = writer.write_markdown(_note(), tmp_path)
    content = path.read_text(encoding="utf-8")
    assert "  - groceries" in content
    assert "  - errand" in content


def test_write_markdown_with_no_tags_writes_an_empty_list(tmp_path, frozen_time):
    path = writer.write_markdown(_note(tags=[]), tmp_path)
    content = path.read_text(encoding="utf-8")
    assert "tags:\n  []" in content


def test_write_markdown_with_no_reminders_writes_an_empty_list(tmp_path, frozen_time):
    path = writer.write_markdown(_note(reminders=[]), tmp_path)
    content = path.read_text(encoding="utf-8")
    assert "reminders: []" in content


def test_write_markdown_includes_scheduled_reminders(tmp_path, frozen_time):
    note = _note(reminders=[{"text": "call the dentist", "datetime": "next Monday 9am"}])
    path = writer.write_markdown(note, tmp_path)
    content = path.read_text(encoding="utf-8")
    assert 'text: "call the dentist"' in content
    assert 'datetime: "next Monday 9am"' in content


def test_write_markdown_escapes_double_quotes_in_summary(tmp_path, frozen_time):
    note = _note(summary='He said "hi" to me.')
    path = writer.write_markdown(note, tmp_path)
    content = path.read_text(encoding="utf-8")
    assert "He said 'hi' to me." in content


def test_write_markdown_preserves_the_raw_transcript_on_a_fallback_note(tmp_path, frozen_time):
    """A fallback note (Ollama unreachable) has body == the raw transcript.
    This is the on-disk half of the graceful-degradation guarantee."""
    transcript = "call the dentist tomorrow at nine"
    note = _note(body=transcript, _fallback_reason="Ollama not running")
    path = writer.write_markdown(note, tmp_path)
    content = path.read_text(encoding="utf-8")
    assert transcript in content


# -- write_docx -------------------------------------------------------------

def test_write_docx_creates_a_document_with_the_title_as_heading(tmp_path, frozen_time):
    from docx import Document

    path = writer.write_docx(_note(), tmp_path)
    assert path.exists()

    doc = Document(str(path))
    assert doc.paragraphs[0].text == "Buy groceries"


def test_write_docx_renders_bullet_list_items_from_the_body(tmp_path, frozen_time):
    from docx import Document

    note = _note(body="Intro line\n- first item\n- second item")
    path = writer.write_docx(note, tmp_path)

    doc = Document(str(path))
    bullet_texts = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "first item" in bullet_texts
    assert "second item" in bullet_texts
